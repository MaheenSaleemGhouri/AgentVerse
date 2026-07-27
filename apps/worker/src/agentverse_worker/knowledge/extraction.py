"""Raw bytes → plain text, per document type.

Kept separate from chunking (which is pure and shared) because extraction
is inherently I/O-shaped and format-library-dependent. Every extractor
returns text only; deciding how to *split* that text is
`agentverse_shared.text.chunking`'s job.

Security posture (`secure-coding-expert`): the caller has already
content-sniffed and size-capped the upload. These functions additionally
refuse to trust a declared type — the extractor is selected by sniffed
signature first, so a `.pdf` that is really a zip does not reach the PDF
parser.
"""

from __future__ import annotations

import csv
import io
import json
import logging
from dataclasses import dataclass

import pypdf
from docx import Document as DocxDocument

logger = logging.getLogger(__name__)

#: Magic-byte signatures we can positively identify. Anything else is
#: treated as text and must decode as UTF-8 to be accepted.
_PDF_MAGIC = b"%PDF-"
_ZIP_MAGIC = b"PK\x03\x04"  # DOCX is a zip container


class ExtractionError(Exception):
    """Raised when content cannot be turned into text. The ingestion job
    turns this into a `failed` document with a user-visible reason rather
    than a crashed worker.
    """


@dataclass(frozen=True, slots=True)
class ExtractedText:
    text: str
    #: What the extractor concluded the bytes actually were, which may
    #: differ from the client's declared content type. Recorded so a
    #: mismatch is visible in logs rather than silently reinterpreted.
    detected_type: str


def extract_text(data: bytes, *, filename: str, declared_content_type: str) -> ExtractedText:
    """Extracts plain text, dispatching on sniffed content, not on the
    client-declared type.
    """
    if not data:
        raise ExtractionError("file is empty")

    if data.startswith(_PDF_MAGIC):
        return ExtractedText(text=_extract_pdf(data), detected_type="application/pdf")

    if data.startswith(_ZIP_MAGIC):
        # Only DOCX is supported among zip containers; a generic zip
        # upload must be rejected rather than parsed hopefully.
        if not filename.lower().endswith((".docx", ".docm")):
            raise ExtractionError("zip-based files are only supported as .docx Word documents")
        return ExtractedText(
            text=_extract_docx(data),
            detected_type=(
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            ),
        )

    text = _decode_text(data)
    lowered = filename.lower()

    if lowered.endswith((".json",)):
        return ExtractedText(text=_flatten_json(text), detected_type="application/json")
    if lowered.endswith((".csv", ".tsv")):
        return ExtractedText(text=_normalize_csv(text, lowered), detected_type="text/csv")

    return ExtractedText(text=text, detected_type=declared_content_type or "text/plain")


def _decode_text(data: bytes) -> str:
    """Strict UTF-8. Silently replacing undecodable bytes would embed
    mojibake and quietly poison retrieval for that document.
    """
    try:
        return data.decode("utf-8")
    except UnicodeDecodeError as exc:
        raise ExtractionError("file is not valid UTF-8 text") from exc


def _extract_pdf(data: bytes) -> str:
    try:
        reader = pypdf.PdfReader(io.BytesIO(data))
        if reader.is_encrypted:
            raise ExtractionError("encrypted PDFs are not supported")
        pages = [(page.extract_text() or "").strip() for page in reader.pages]
    except ExtractionError:
        raise
    except Exception as exc:
        raise ExtractionError(f"could not read PDF: {exc}") from exc

    text = "\n\n".join(p for p in pages if p)
    if not text.strip():
        # A scanned/image-only PDF extracts to nothing. Failing loudly
        # beats indexing an empty document the user believes is searchable.
        raise ExtractionError(
            "no extractable text found — the PDF may be scanned images (OCR is not supported)"
        )
    return text


def _extract_docx(data: bytes) -> str:
    try:
        document = DocxDocument(io.BytesIO(data))
        blocks = [p.text.strip() for p in document.paragraphs]
        for table in document.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                if any(cells):
                    blocks.append(" | ".join(cells))
    except Exception as exc:
        raise ExtractionError(f"could not read Word document: {exc}") from exc

    text = "\n\n".join(b for b in blocks if b)
    if not text.strip():
        raise ExtractionError("no extractable text found in the Word document")
    return text


def _flatten_json(text: str) -> str:
    """Renders JSON as one record per line.

    Line-per-record is what the `structured` chunker expects, so a JSON
    array of objects chunks record-wise instead of being split mid-object
    by a pretty-printer's arbitrary newlines.
    """
    try:
        parsed = json.loads(text)
    except json.JSONDecodeError as exc:
        raise ExtractionError(f"invalid JSON: {exc}") from exc

    if isinstance(parsed, list):
        return "\n".join(json.dumps(item, ensure_ascii=False) for item in parsed)
    return json.dumps(parsed, ensure_ascii=False)


def _normalize_csv(text: str, lowered_filename: str) -> str:
    """Re-emits CSV/TSV as comma-delimited lines with quoting resolved.

    Parsing then re-emitting collapses embedded newlines inside quoted
    fields onto one line, so the `structured` chunker's one-line-per-record
    assumption actually holds.
    """
    delimiter = "\t" if lowered_filename.endswith(".tsv") else ","
    try:
        rows = list(csv.reader(io.StringIO(text), delimiter=delimiter))
    except csv.Error as exc:
        raise ExtractionError(f"could not parse delimited file: {exc}") from exc

    out = io.StringIO()
    writer = csv.writer(out, lineterminator="\n")
    for row in rows:
        if any(cell.strip() for cell in row):
            writer.writerow([cell.replace("\n", " ").strip() for cell in row])
    return out.getvalue().rstrip("\n")
