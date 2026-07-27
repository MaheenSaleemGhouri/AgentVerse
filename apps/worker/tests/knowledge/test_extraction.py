"""Extraction tests build real files (a real DOCX via python-docx, a real
PDF via pypdf) rather than asserting against hand-faked bytes, so a
library upgrade that changes parsing behavior actually shows up here.
"""

from __future__ import annotations

import io

import pypdf
import pytest
from docx import Document as DocxDocument

from agentverse_worker.knowledge.extraction import ExtractionError, extract_text


def _docx_bytes(paragraphs: list[str], table: list[list[str]] | None = None) -> bytes:
    document = DocxDocument()
    for text in paragraphs:
        document.add_paragraph(text)
    if table:
        docx_table = document.add_table(rows=len(table), cols=len(table[0]))
        for r, row in enumerate(table):
            for c, cell in enumerate(row):
                docx_table.rows[r].cells[c].text = cell
    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def _blank_pdf_bytes() -> bytes:
    writer = pypdf.PdfWriter()
    writer.add_blank_page(width=200, height=200)
    buffer = io.BytesIO()
    writer.write(buffer)
    return buffer.getvalue()


# --- plain text ---------------------------------------------------------


def test_extracts_plain_text() -> None:
    result = extract_text(
        b"Hello world.\n\nSecond paragraph.",
        filename="notes.txt",
        declared_content_type="text/plain",
    )

    assert "Hello world." in result.text
    assert "Second paragraph." in result.text


def test_rejects_empty_file() -> None:
    with pytest.raises(ExtractionError, match="empty"):
        extract_text(b"", filename="a.txt", declared_content_type="text/plain")


def test_rejects_invalid_utf8_rather_than_producing_mojibake() -> None:
    # Replacing undecodable bytes would silently embed garbage and
    # degrade retrieval for that document with no visible error.
    with pytest.raises(ExtractionError, match="not valid UTF-8"):
        extract_text(b"\xff\xfe\x00\x01", filename="a.txt", declared_content_type="text/plain")


# --- DOCX ---------------------------------------------------------------


def test_extracts_docx_paragraphs() -> None:
    data = _docx_bytes(["First paragraph.", "Second paragraph."])

    result = extract_text(
        data,
        filename="memo.docx",
        declared_content_type=(
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        ),
    )

    assert "First paragraph." in result.text
    assert "Second paragraph." in result.text


def test_extracts_docx_tables() -> None:
    # Tables often carry the actual facts a user wants retrievable;
    # paragraph-only extraction would silently drop them.
    data = _docx_bytes(["Intro."], table=[["Region", "Revenue"], ["EMEA", "1200"]])

    result = extract_text(data, filename="report.docx", declared_content_type="")

    assert "Region | Revenue" in result.text
    assert "EMEA | 1200" in result.text


def test_empty_docx_is_rejected() -> None:
    data = _docx_bytes([])

    with pytest.raises(ExtractionError, match="no extractable text"):
        extract_text(data, filename="blank.docx", declared_content_type="")


def test_non_docx_zip_is_rejected() -> None:
    # A .zip renamed to .docx must not be hopefully parsed.
    import zipfile

    buffer = io.BytesIO()
    with zipfile.ZipFile(buffer, "w") as zf:
        zf.writestr("payload.txt", "not a word document")

    with pytest.raises(ExtractionError, match="only supported as .docx"):
        extract_text(buffer.getvalue(), filename="archive.zip", declared_content_type="")


# --- PDF ----------------------------------------------------------------


def test_pdf_with_no_extractable_text_is_rejected() -> None:
    # A scanned/image-only PDF yields no text. Failing loudly beats
    # indexing an empty document the user believes is searchable.
    with pytest.raises(ExtractionError, match="no extractable text"):
        extract_text(
            _blank_pdf_bytes(), filename="scan.pdf", declared_content_type="application/pdf"
        )


def test_corrupt_pdf_is_rejected_with_a_readable_reason() -> None:
    with pytest.raises(ExtractionError, match="could not read PDF"):
        extract_text(
            b"%PDF-1.4 this is not actually a valid pdf body",
            filename="broken.pdf",
            declared_content_type="application/pdf",
        )


# --- sniffing beats the declared type -----------------------------------


def test_a_pdf_declared_as_text_is_still_parsed_as_a_pdf() -> None:
    # The client-declared content type is untrusted; magic bytes decide.
    with pytest.raises(ExtractionError, match="no extractable text"):
        extract_text(_blank_pdf_bytes(), filename="mystery.txt", declared_content_type="text/plain")


def test_text_declared_as_pdf_is_not_sent_to_the_pdf_parser() -> None:
    result = extract_text(
        b"Actually just prose.", filename="mislabeled.pdf", declared_content_type="application/pdf"
    )

    assert "Actually just prose." in result.text


# --- structured normalization ------------------------------------------


def test_json_array_becomes_one_record_per_line() -> None:
    # The structured chunker packs line-per-record; a pretty-printed blob
    # would otherwise be split mid-object.
    data = b'[{"a": 1}, {"a": 2}, {"a": 3}]'

    result = extract_text(data, filename="items.json", declared_content_type="application/json")

    assert len(result.text.splitlines()) == 3


def test_invalid_json_is_rejected() -> None:
    with pytest.raises(ExtractionError, match="invalid JSON"):
        extract_text(b"{not json", filename="bad.json", declared_content_type="application/json")


def test_csv_quoted_newlines_are_flattened_onto_one_line() -> None:
    # Without this, one record spans several lines and the "one line per
    # record" chunking assumption silently breaks.
    data = b'name,note\nalpha,"line one\nline two"\nbeta,ok'

    result = extract_text(data, filename="data.csv", declared_content_type="text/csv")

    assert len(result.text.splitlines()) == 3


def test_tsv_is_parsed_with_tab_delimiter() -> None:
    data = b"name\tvalue\nalpha\t1"

    result = extract_text(
        data, filename="data.tsv", declared_content_type="text/tab-separated-values"
    )

    assert "name,value" in result.text
